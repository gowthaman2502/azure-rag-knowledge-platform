from io import BytesIO
from pathlib import PurePosixPath

import pandas as pd
from azure.storage.blob import BlobServiceClient
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument


from rag_knowledge_assistant.config import AZURE_STORAGE_CONNECTION_STRING, AZURE_STORAGE_CONTAINER


class KnowledgeBaseIngestion:

    def __init__(self):
        self.blob_service_client = BlobServiceClient.from_connection_string(
            AZURE_STORAGE_CONNECTION_STRING
        )
        self.container_client = self.blob_service_client.get_container_client(
            AZURE_STORAGE_CONTAINER
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1200,
            chunk_overlap=150,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    def list_documents(self):
        supported = {".pdf", ".docx", ".xlsx"}
        return [
            blob.name
            for blob in self.container_client.list_blobs()
            if PurePosixPath(blob.name).suffix.lower() in supported
        ]

    def download_document(self, blob_name):
        blob_client = self.container_client.get_blob_client(blob_name)
        return blob_client.download_blob().readall()

    def load_document(self, blob_name):
        file_bytes = self.download_document(blob_name)
        extension = PurePosixPath(blob_name).suffix.lower()
        department = PurePosixPath(blob_name).parts[0]

        if extension == ".pdf":
            documents = self._load_pdf(file_bytes)
        elif extension == ".docx":
            documents = self._load_docx(file_bytes)
        elif extension == ".xlsx":
            documents = self._load_xlsx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        for document in documents:
            document.metadata.update({
                "document_name": PurePosixPath(blob_name).name,
                "document_path": blob_name,
                "document_type": extension.replace(".", ""),
                "department": department
            })

        return documents

    def _load_pdf(self, file_bytes):
        reader = PdfReader(BytesIO(file_bytes))
        documents = []

        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()

            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={"page": page_number}
                    )
                )

        return documents

    def _load_docx(self, file_bytes):
        document = DocxDocument(BytesIO(file_bytes))
        sections = []
        current_section = "General"
        current_text = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if not text:
                continue

            style = paragraph.style.name.lower() if paragraph.style else ""

            is_heading = "heading" in style

            if is_heading and current_text:
                sections.append(
                    Document(
                        page_content="\n".join(current_text),
                        metadata={"section": current_section}
                    )
                )
                current_text = []

            if is_heading:
                current_section = text

            current_text.append(text)

        if current_text:
            sections.append(
                Document(
                    page_content="\n".join(current_text),
                    metadata={"section": current_section}
                )
            )

        return sections

    def _load_xlsx(self, file_bytes):
        excel = pd.ExcelFile(BytesIO(file_bytes))
        documents = []

        for sheet_name in excel.sheet_names:
            dataframe = pd.read_excel(excel, sheet_name=sheet_name)

            text = f"Sheet: {sheet_name}\n{dataframe.to_csv(index=False)}"

            documents.append(
                Document(
                    page_content=text,
                    metadata={"sheet_name": sheet_name}
                )
            )

        return documents

    def create_chunks(self, documents):
        chunks = []

        for document in documents:
            split_documents = self.text_splitter.split_documents([document])
            chunks.extend(split_documents)

        counters = {}

        for chunk in chunks:
            document_name = chunk.metadata.get("document_name", "document")
            chunk_id = counters.get(document_name, 0)
            counters[document_name] = chunk_id + 1
            chunk.metadata["chunk_id"] = chunk_id

        return chunks

    def run(self):
        documents = self.list_documents()
        print(f"Found {len(documents)} documents")

        all_chunks = []

        for blob_name in documents:
            print(f"Processing: {blob_name}")

            loaded_documents = self.load_document(blob_name)
            chunks = self.create_chunks(loaded_documents)
            all_chunks.extend(chunks)

            print(
                f"  Loaded: {len(loaded_documents)} "
                f"chunks: {len(chunks)}"
            )

        print(f"\nTotal chunks: {len(all_chunks)}")
        return all_chunks